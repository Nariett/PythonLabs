#код запускать двойным нажатием
import tkinter as tk
import tkinter.messagebox
from tkinter import ttk
from tkinter import *
import sqlite3
from datetime import datetime
import docx 
from docx import document

path = "C:/Users/Sasha/Desktop/Shop/ShopDB.db"#путь к бд

def is_numeric_string(s):
    try:
        float(s)
        return True
    except ValueError:
        return False

def add_product_form(event):
    new_window = tk.Toplevel(root)
    new_window.title("Добавить товар")    
    new_window.geometry("300x100")

    product_name_label = tk.Label(new_window, text="Название:")
    product_name_label.grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
    product_name_entry = tk.Entry(new_window)
    product_name_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)

    product_price_label = tk.Label(new_window, text="Цена:")
    product_price_label.grid(row=1, column=0, padx=5, pady=5, sticky=tk.W)
    product_price_entry = tk.Entry(new_window)
    product_price_entry.grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)

    add_product_button = tk.Button(new_window, text="Добавить", command=lambda: add_product(new_window, product_name_entry, product_price_entry))
    add_product_button.grid(row=2, column=0, columnspan=2, padx=5, pady=5)


def add_product(new_window, name_entry, price_entry):
    name = name_entry.get()
    string = price_entry.get()
    price = None 
    if is_numeric_string(string):
        price = float(string)
    else:
        new_window.destroy()
        tkinter.messagebox.showinfo("Ошибка", "Введите корректную цену" )
    if (name is None or price is None):
        new_window.destroy()
        tkinter.messagebox.showinfo("Ошибка", "Товар не добавлен. Введите название и цену" )
    elif (price < float(0)):
        new_window.destroy()
        tkinter.messagebox.showinfo("Ошибка", "Товар не добавлен. Введите цену больше нуля" )
    else:
        con = sqlite3.connect(path)
        cursor = con.cursor()
        cursor.execute("INSERT INTO Product (Name, Price) VALUES (?, ?)", (name, string))
        con.commit()  
        con.close()  
        new_window.destroy()
        tkinter.messagebox.showinfo("Уведомление", "Товар добавлен")
        update_data(root)

    
def add_customer_form(event):
    new_window = tk.Toplevel(root)
    new_window.title("Добавить клиента")   
    new_window.geometry("300x100")

    customer_name_label = tk.Label(new_window, text="ФИО:")
    customer_name_label.grid(row=0, column=0, padx=5, pady=5,  sticky=tk.W)
    customer_name_entry = tk.Entry(new_window)
    customer_name_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)

    customer_number_label = tk.Label(new_window, text="Номер телефона:")
    customer_number_label.grid(row=1, column=0, padx=5, pady=5, sticky=tk.W)
    customer_number_entry = tk.Entry(new_window)
    customer_number_entry.grid(row=1, column=1, padx=5, pady=5,  sticky=tk.W)

    add_product_button = tk.Button(new_window, text="Добавить", command=lambda: add_customer(new_window,customer_name_entry, customer_number_entry))
    add_product_button.grid(row=2, column=0, columnspan=2, padx=5, pady=5)

def add_customer(new_window, name_entry, number_entry):
    name = name_entry.get()
    number = number_entry.get()
    if (name is None or number is None):
        new_window.destroy()
        tkinter.messagebox.showinfo("Ошибка", "Клиент не добавлен. Введите ФИО и номер телефона" )
    elif (len(number) < 5):
        new_window.destroy()
        tkinter.messagebox.showinfo("Ошибка", "Клиент не добавлен. Введите корректный номер телефона" )
    else:
        con = sqlite3.connect(path)
        cursor = con.cursor()
        cursor.execute("INSERT INTO Customer (Name, PhoneNumber) VALUES (?, ?)", (name, number))
        con.commit()  
        con.close()  
        tkinter.messagebox.showinfo("Уведомление", "Клиент добавлен")
        new_window.destroy()
        update_data(root)


def add_all(new_window, customer, product,quantity,price):
    if (len(customer) == 0 or len(product) == 0 or len(quantity) == 0 or len(price) == 0):
        tkinter.messagebox.showinfo("Ошибка", "Заказ не добавлен. Заполните все поля" )
    else:
        price = float(price)
        current_time = f"{datetime.now():%d-%m-%Y}"
        conn = sqlite3.connect(path)
        cur = conn.cursor()
        cur.execute("SELECT ID_Client FROM Customer WHERE Name = ?", (customer,))
        result = cur.fetchone()
        id_customer = result[0]
        cur.execute("SELECT ID_Product FROM Product WHERE Name = ?", (product,))
        result2 = cur.fetchone()
        id_product = result2[0]
        cur.execute("INSERT INTO Order_List (ID_Customer, Date) VALUES (?,?)", (int(id_customer),str(current_time),))
        conn.commit()  
        cur.execute("INSERT INTO Shopping_List (ID_Customer,ID_Product,Quantity,Price) VALUES (?,?,?,?)",(int(id_customer),int(id_product),int(quantity),float(price),))
        conn.commit()
        conn.close()  
        tkinter.messagebox.showinfo("Уведомление", "Заказ добавлен")

def ShowCheck(customer):
    
    connect = sqlite3.connect(path)
    cursor = connect.cursor()
    cursor.execute("SELECT ID_Client FROM Customer WHERE Name = ?", (customer,))
    index = cursor.fetchone()
    cursor.execute("SELECT ID_Product, Quantity, Price FROM Shopping_List WHERE Id_Customer = ?", (index[0],))
    list_Order = cursor.fetchall()
    document = docx.Document()
    document.add_heading('Чек', 0)
    document.add_paragraph('Клиент: ' + str(customer))
    FullPrice = 0
    for row in list_Order:
        cursor.execute("SELECT Name FROM Product WHERE ID_Product = ?", (int(row[0]),))
        Name_Product = cursor.fetchone()
        document.add_paragraph('Товар: '+ str(Name_Product[0]))
        document.add_paragraph('Кол-во: ' + str(row[1]))
        document.add_paragraph('К оплате: ' + str(row[2]) + "₽")
        FullPrice+=row[2]
    document.add_paragraph('Сумма общего чека: ' + str(FullPrice) + "₽")
    document.save("Чек.docx")
    tkinter.messagebox.showinfo("Готово", "Чек создан.")

#global переменные 
listProduct = []
listCustomer = []
listPrice = []


def spinbox_updated(*args):
    global listProduct
    global combo
    global listPrice
    value = spinbox_value.get()
    item = combo2.get()
    index = listProduct.index(item) 
    Price = listPrice[index]
    label_one.config(text=str(int(value )* Price ))

def update_data(root):
    global listProduct
    global listCustomer
    global listPrice
    conn = sqlite3.connect(path)
    cur = conn.cursor()
    cur.execute("SELECT Name FROM Customer")
    data = cur.fetchall()
    values = [row[0] for row in data]
    cur.close()
    conn.close()
    listCustomer = values
    conn = sqlite3.connect(path)
    cur = conn.cursor()
    cur.execute("SELECT Name FROM Product")
    data = cur.fetchall()
    values2=[row[0] for row in data]
    cur.close()
    conn.close()
    listProduct = values2
    conn = sqlite3.connect(path)
    cur = conn.cursor()
    cur.execute("SELECT Price FROM Product")
    data = cur.fetchall()
    valuesPrice=[row[0] for row in data]
    cur.close()
    conn.close()
    listPrice = valuesPrice
    create_widgets(root)

def create_widgets(root):
    global combo
    global combo2
    combo = ttk.Combobox(root, values=listCustomer, width=10)
    combo.grid(row=1, column=1, columnspan=2, padx=10, pady=10,sticky=tk.W)# 1 1 
    combo2 = ttk.Combobox(root, values=listProduct, width=10)
    combo2.grid(row=1, column=4, columnspan=2, padx=10, pady=10,sticky=tk.W)# 1 4 

def update_combobox():
    global combo
    global combo2
    combo.config(values=listCustomer)
    combo2.config(values=listProduct)


root = tk.Tk()
update_data(root)
create_widgets(root)
root.title("Главная форма")
root.geometry("700x300+0+0")
button_one = tk.Button(root, text="Добавить Товар")
button_one.grid(row=0, column=0, pady=10, padx=10)
button_one.bind("<Button-1>", add_product_form)
button_two = tk.Button(root, text="Добавить Клиента")
button_two.grid(row=0, column=1, pady=10, padx=10, sticky="w")
button_two.bind("<Button-1>", add_customer_form)
tk.Label(text="Клиент:").grid(row=1, column=0, sticky=tk.W, padx=10, pady=10)
tk.Label(text="Товар:").grid(row=1, column=3, sticky=tk.W, padx=10, pady=10)
tk.Label(text="Кол-во:").grid(row=1, column=6, sticky=tk.W, padx=10, pady=10)
spinbox_value = tk.StringVar()
spinbox = tk.Spinbox(root, width=7, from_=1, to=50, textvariable=spinbox_value)
spinbox.grid(row=1, column=7, padx=10,sticky=tk.W) 
spinbox_value.trace('w', lambda *args: spinbox_updated(*args))
tk.Label(text="Сумма:").grid(row=1, column=8, sticky=tk.W, padx=10, pady=10)
label_one = tk.Label(text="")
label_one.grid(row=1, column=9, sticky=tk.W, padx=10, pady=10)
add_product_button = tk.Button(root, text="Оформить заказ", command=lambda: add_all(root,combo.get(),combo2.get(),spinbox.get(), label_one.cget("text")))
add_product_button.grid(row=5, column=0, columnspan=1, padx=10, pady=10,sticky=tk.W) 
print_order_button = tk.Button(root, text="Напечатать чек", command=lambda: ShowCheck(combo.get()))##ShowCheck(combo.get()
print_order_button.grid(row=5, column=1, columnspan=1, padx=10, pady=10,sticky=tk.W) 

root.mainloop()
